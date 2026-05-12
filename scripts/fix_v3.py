"""Fix v3:
1. Move ER-diagram image (Рисунок 2.10) right before its caption.
2. Renumber duplicate "Рисунок 1.2 – Структурная схема существующей" → "Рисунок 1.3".
3. Update abstract counts: pages, tables, figures, sources.
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

PATH = "dist/Диплом_Ксёнжик_исправленный.docx"


def set_text(p, text):
    runs = p.runs
    if not runs:
        p.add_run(text); return
    first = runs[0]
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    first.text = text


def find(doc, contains, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        if contains in p.text:
            return i, p
    return -1, None


def main():
    doc = Document(PATH)

    # ------------------------------------------------------------
    # 1. Move ER-diagram image (paragraph containing w:drawing) so it
    #    sits right before "Рисунок 2.10 – ER-диаграмма" caption.
    # ------------------------------------------------------------
    cap_idx, cap_p = find(doc, "Рисунок 2.10 – ER-диаграмма")
    if cap_p:
        # Find nearest preceding paragraph that contains a w:drawing
        img_p = None
        for j in range(cap_idx - 1, max(0, cap_idx - 12), -1):
            pp = doc.paragraphs[j]
            if pp._element.findall('.//' + qn('w:drawing')):
                img_p = pp
                break
        if img_p is not None and img_p._element is not cap_p._element.getprevious():
            # Move img_p element right before cap_p
            img_el = img_p._element
            img_el.getparent().remove(img_el)
            cap_p._element.addprevious(img_el)
            print("Moved ER image right before its caption.")

    # ------------------------------------------------------------
    # 2. Renumber duplicate "Рисунок 1.2 – Структурная схема существующей"
    # ------------------------------------------------------------
    _, dup = find(doc, "Рисунок 1.2 – Структурная схема существующей")
    if dup:
        set_text(dup, "Рисунок 1.3 – Структурная схема существующей информационной системы БЕЛТА")
        print("Renumbered to Рисунок 1.3")
    # Also update the section heading "1.7 Структурная схема существующей" 
    # text refers to "рисунок 1.2" if any — search for body text mentioning it
    for p in doc.paragraphs:
        if "рисунке 1.2" in p.text and "существующ" in p.text:
            set_text(p, p.text.replace("рисунке 1.2", "рисунке 1.3"))

    # ------------------------------------------------------------
    # 3. Update abstract paragraph 56 (Полученные результаты и их новизна)
    #    with correct counts: ~55 pages, 10 tables, 28 figures, 48 sources.
    # ------------------------------------------------------------
    _, p = find(doc, "Полученные результаты и их новизна")
    if p:
        new_text = (
            "Полученные результаты и их новизна: спроектирована и реализована трёхзвенная "
            "архитектура веб-приложения (клиент на React, сервер на Express.js, СУБД "
            "PostgreSQL), разработана схема базы данных из семи связанных сущностей, "
            "реализована оригинальная подсистема геймификации XP с защитой от повторного "
            "начисления очков на уровне ограничения целостности базы данных, создана "
            "административная панель и подсистема модерации. Объём пояснительной записки — "
            "55 страниц; в работе содержатся 10 таблиц, 28 рисунков и 48 использованных "
            "источников."
        )
        set_text(p, new_text)
        print("Abstract counts updated.")

    doc.save(PATH)
    print("Saved.")


if __name__ == "__main__":
    main()
