"""Apply v2 supervisor corrections to the already-edited thesis:
- Insert real diagrams under Рисунок 1.1, Рисунок 1.2, Рисунок 2.9а
- Rewrite section 2.8 (was about e-commerce, must be about news)
"""
from copy import deepcopy
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

SRC = "dist/Диплом_Ксёнжик_исправленный.docx"
DST = "dist/Диплом_Ксёнжик_исправленный.docx"


def set_text(p, text):
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    first = runs[0]
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    first.text = text


def insert_image_before(caption_p, image_path, width_cm=15.5):
    """Insert a new centered image paragraph immediately before caption_p."""
    # Clone caption paragraph element to preserve style/alignment, then replace its
    # contents with an inline image.
    new_el = deepcopy(caption_p._element)
    caption_p._element.addprevious(new_el)
    new_p = Paragraph(new_el, caption_p._parent)
    # Clear text in all runs
    for r in list(new_p.runs):
        r._element.getparent().remove(r._element)
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    return new_p


def find(doc, contains, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if contains in p.text:
            return i, p
    return -1, None


def main():
    doc = Document(SRC)

    # ============================================================
    # 1. Insert image under "Рисунок 1.1 – Организационная структура РУП «БЕЛТА»"
    # ============================================================
    _, p = find(doc, "Рисунок 1.1 – Организационная структура")
    if p:
        insert_image_before(p, "dist/figures/fig_1_1_belta_org.png", width_cm=15.5)

    # ============================================================
    # 2. Insert image under "Рисунок 1.2 – Структурная схема существующей"
    # ============================================================
    _, p = find(doc, "Рисунок 1.2 – Структурная схема существующей")
    if p:
        insert_image_before(p, "dist/figures/fig_1_2_existing_system.png", width_cm=15.5)

    # ============================================================
    # 3. Insert image under "Рисунок 2.9а – Структурная схема разработанной"
    # ============================================================
    _, p = find(doc, "Рисунок 2.9а – Структурная схема разработанной")
    if p:
        insert_image_before(p, "dist/figures/fig_2_9a_developed_system.png", width_cm=16.0)

    # ============================================================
    # 4. Rewrite section 2.8 — was about e-commerce (товары, корзина, заказ)
    # ============================================================
    # 2.8 paragraphs found at 522 (heading), 523 (body), 524 (table caption),
    # 525 (note), 526 (body), 527 (body), 528 (next section heading).
    fixes_28 = {
        "Проверка пользовательских функций охватывает сценарии, с которых начинается":
            "Проверка пользовательских функций охватывает основные сценарии работы посетителя "
            "и зарегистрированного пользователя новостной платформы: регистрацию учётной записи "
            "и вход в систему, просмотр ленты новостей и страницы отдельной статьи, поиск "
            "материалов по ключевым словам, фильтрацию по тематическим категориям, оставление "
            "и удаление комментариев, добавление публикаций в закладки и работу с личным "
            "кабинетом, включая отображение текущего количества опытных очков (XP), уровня "
            "пользователя и истории активности.",
        "Тест-кейсы данной группы подтверждают, что пользовательский путь реализован":
            "Тест-кейсы данной группы подтверждают, что пользовательский путь реализован "
            "последовательно и без разрывов между страницами. Особое внимание уделяется "
            "операциям, связанным с начислением опытных очков и переходом пользователя на "
            "следующий уровень, а также корректности работы закладок и комментариев, "
            "поскольку именно эти функции формируют интерактивную часть новостной платформы "
            "и влияют на удержание аудитории.",
    }
    for needle, new_text in fixes_28.items():
        _, pp = find(doc, needle)
        if pp:
            set_text(pp, new_text)

    # The table caption at para 524 is fine ("Тест-кейсы основных пользовательских функций"),
    # but if it still mentions товары, fix it.
    _, pp = find(doc, "Таблица 2.1")
    if pp and ("товар" in pp.text.lower() or "корзин" in pp.text.lower()):
        set_text(pp, "Таблица 2.1 — Тест-кейсы основных пользовательских функций")

    doc.save(DST)
    print("Saved:", DST)


if __name__ == "__main__":
    main()
